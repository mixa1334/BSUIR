import docx
import pickle

import Vocabulary
import Sentence

import nltk
from nltk.corpus import wordnet as wn
from string import punctuation


def read_text_from_file(filename):
    doc = docx.Document(filename)
    text = []
    for paragraph in doc.paragraphs:
        text.append(paragraph.text)
    return "\n".join(text)


def write_text_to_file(filename, voc):
    doc_file = docx.Document()
    doc_file.add_heading("Предложения", 1)
    for sentence in dict(voc.get_all_sentences()).values():
        doc_file.add_paragraph("---------------------------------------------------------------------------------")
        sentence_text = "Предложение: " + "\'" + str(sentence.get_string()) + "\"" + "\n дерево -> " + str(
            sentence.get_tree())
        doc_file.add_paragraph(sentence_text)
    doc_file.save(filename + ".docx")


def write_vocabulary_to_file(voc, filename):
    with open(filename + ".pkl", "wb") as file:
        result_list = []
        for sent in dict(voc.get_all_sentences()).values():
            sent_list = [sent.get_string(), sent.get_elements(), sent.get_parse_result()]
            result_list.append(sent_list)
        pickle.dump(result_list, file)


def read_vocabulary_from_file(filename):
    with open(filename, "rb") as file:
        voc = Vocabulary.Vocabulary()
        result_list = list(pickle.load(file))
        for item in result_list:
            s = Sentence.Sentence()
            s.set_string(item[0])
            s.set_elements(list(item[1]))
            parse_result = str(item[2])
            s.set_parse_result(parse_result)
            s.set_tree(nltk.tree.Tree.fromstring(parse_result))
            voc.add_sentence(s)
        return voc


def process_text(text):
    voc = Vocabulary.Vocabulary()

    sentence_list = list(nltk.tokenize.sent_tokenize(text))
    for sent in sentence_list:
        s = process_sentence(sent)
        voc.add_sentence(s)

    return voc


def process_sentence(sentence):
    s = Sentence.Sentence()
    tokens = list(filter(lambda val: val not in punctuation, nltk.word_tokenize(sentence)))

    result = '(S '
    for token in tokens:
        result += get_word_info(token)
    result += ')'
    tree = nltk.tree.Tree.fromstring(result)

    s.set_string(sentence)
    s.set_elements(tokens)
    s.set_tree(tree)
    s.set_parse_result(result)

    return s


def get_word_info(word):
    if len(wn.synsets(word)) == 0:
        return '(WS (W ' + word + '))'
    result = '(WS (W ' + word + ') (DEF ' + \
             wn.synsets(word)[0].definition().replace(' ', '_') + ')'
    synonyms, antonyms, hyponyms, hypernyms = [], [], [], []
    word = wn.synsets(word)
    print(word[0])
    syn_app = synonyms.append
    ant_app = antonyms.append
    he_app = hyponyms.append
    hy_app = hypernyms.append
    for synset in word:
        for lemma in synset.lemmas():
            syn_app(lemma.name())
            if lemma.antonyms():
                ant_app(lemma.antonyms()[0].name())
    for hyponym in word[0].hyponyms():
        he_app(hyponym.name())
    for hypernym in word[0].hypernyms():
        hy_app(hypernym.name())
    if len(synonyms):
        result += ' (SYN '
        for synonym in synonyms:
            result += synonym + ' '
    if len(antonyms):
        result += ') (ANT '
        for antonym in antonyms:
            result += antonym + ' '
    if len(hyponyms):
        result += ') (HY '
        for hyponym in hyponyms:
            result += hyponym + ' '
    if len(hypernyms):
        result += ') (HE '
        for hypernym in hypernyms:
            result += hypernym + ' '
    result += '))'
    return result
